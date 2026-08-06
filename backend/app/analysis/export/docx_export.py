"""Word (.docx) report export.

Produces the document an officer files or hands over: the authority's heading, the
case identifiers, the assessment and its basis, then the supporting evidence in
full. The deterministic verdict from `summary.py` is always what the document
states — the narrative writer only supplies prose around it, and where that prose
came from is disclosed in the document itself.
"""

from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from ..narrative import SECTIONS, NarrativeResult

VERDICT_STATEMENTS = {
    "malicious": "MALICIOUS",
    "suspicious": "SUSPICIOUS — REQUIRES FURTHER EXAMINATION",
    "unknown": "NO KNOWN INDICATORS OF COMPROMISE FOUND",
}

NOT_EXECUTED = (
    "The file was not run at any point during this analysis. Every capability "
    "described in this report was determined by reading the file's code. It "
    "describes what the file is able to do, not behaviour that was observed."
)

# Said in place of NOT_EXECUTED once a sample has actually been detonated. The
# distinction is the whole point of the dynamic section: a capability is an
# inference drawn from code, an observation is a record of something that
# happened, and a document that presents one as the other is a document that
# falls apart the first time it is challenged.
EXECUTED = (
    "The file was run on {when} inside a contained sandbox: {containment} The "
    "capabilities listed in this report were still determined by reading the "
    "file's code and describe what it is able to do. Only the entries under "
    "OBSERVED BEHAVIOUR record what the file was seen doing."
)

CONTAINMENT = {
    "android": (
        "an isolated virtual phone, created fresh for this examination, holding "
        "no real accounts, contacts, messages or personal data, and discarded "
        "afterwards."
    ),
    "windows": (
        "an emulated computer, in which the file's instructions were carried out "
        "by a simulated processor against a simulated copy of Windows. The file's "
        "code did not run on a real computer at any point."
    ),
}

DETONATION_ATTEMPTED = (
    "Running the file in a sandbox was attempted but did not complete, so no "
    "behaviour was observed and none is reported. The reason recorded was: {error}"
)

NOT_PROOF_OF_SAFETY = (
    "Where no indicators were found, this is not proof that the file is harmless. "
    "It may be too recent to have been reported, or targeted narrowly enough that "
    "no prior sample has been submitted anywhere."
)


def report_to_docx(
    report: dict[str, Any],
    authority: str,
    narrative: NarrativeResult | None = None,
) -> bytes:
    document = Document()
    _configure_styles(document)
    _letterhead(document, authority)
    _case_details(document, report)
    _verdict_block(document, report)
    _narrative_block(document, report, narrative)
    _evidence(document, report)
    _closing(document, report, narrative)
    _footer(document, report)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# Layout ---------------------------------------------------------------------


def _configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)


def _letterhead(document: Document, authority: str) -> None:
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run(authority.upper())
    run.bold = True
    run.font.size = Pt(17)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Static Malware Analysis Report")
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    _rule(subtitle)
    document.add_paragraph()


def _case_details(document: Document, report: dict[str, Any]) -> None:
    _section_heading(document, "Case details")

    rows = [
        ("File name", report.get("filename") or "—"),
        ("File type", report.get("detected_type") or "—"),
        ("File size", f"{report.get('size', 0):,} bytes"),
        ("SHA-256", report.get("sha256") or "—"),
        ("SHA-1", report.get("sha1") or "—"),
        ("MD5", report.get("md5") or "—"),
        ("Analysed", _stamp(report.get("completed_at") or report.get("created_at"))),
        ("Report generated", _stamp(datetime.now(timezone.utc).isoformat())),
    ]

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for label, value in rows:
        cells = table.add_row().cells
        _cell(cells[0], label, bold=True, width=Pt(120))
        # Hashes are long and must not be re-wrapped mid-string when copied.
        _cell(cells[1], value, mono=label in {"SHA-256", "SHA-1", "MD5"})
    document.add_paragraph()


def _verdict_block(document: Document, report: dict[str, Any]) -> None:
    _section_heading(document, "Assessment")

    verdict = report.get("verdict") or "unknown"
    paragraph = document.add_paragraph()
    run = paragraph.add_run(VERDICT_STATEMENTS.get(verdict, verdict.upper()))
    run.bold = True
    run.font.size = Pt(14)

    if report.get("headline"):
        document.add_paragraph(report["headline"])
    document.add_paragraph()


def _narrative_block(
    document: Document,
    report: dict[str, Any],
    narrative: NarrativeResult | None,
) -> None:
    if narrative and narrative.status == "ok" and narrative.sections:
        for key, title in SECTIONS:
            body = narrative.sections.get(key)
            if not body:
                continue
            _section_heading(document, title)
            for block in body.split("\n"):
                if block.strip():
                    document.add_paragraph(block.strip())
            document.add_paragraph()
        return

    # No narrative writer available: fall back to the summary the analysers
    # produced. The report is never blocked on the model being reachable.
    _section_heading(document, "Summary")
    for block in (report.get("narrative") or "").split("\n\n"):
        if block.strip():
            document.add_paragraph(block.strip())
    document.add_paragraph()


def _evidence(document: Document, report: dict[str, Any]) -> None:
    reasons = report.get("reasons") or []
    if reasons:
        _section_heading(document, "Basis for this assessment")
        for reason in reasons:
            document.add_paragraph(reason, style="List Bullet")
        document.add_paragraph()

    _android(document, report)
    _observed_behaviour(document, report)
    _data_sent_out(document, report)

    techniques = report.get("techniques") or []
    if techniques:
        _section_heading(document, "Capabilities found in the code")
        table = _table(document, ["Capability", "Reference", "Found in code as"])
        for technique in techniques:
            cells = table.add_row().cells
            _cell(cells[0], technique.get("plain_language") or "")
            _cell(cells[1], technique.get("technique_id") or "")
            _cell(cells[2], ", ".join(technique.get("evidence") or []), mono=True)
        document.add_paragraph(
            "These are derived from the program's import table. They describe what "
            "the file is capable of, not actions that were observed."
        )
        document.add_paragraph()

    indicators = report.get("indicators") or []
    if indicators:
        _section_heading(document, "Network destinations referenced by the file")
        table = _table(document, ["Type", "Destination", "Intelligence"])
        for indicator in indicators:
            cells = table.add_row().cells
            _cell(cells[0], indicator.get("type") or "")
            _cell(cells[1], indicator.get("value") or "", mono=True)
            _cell(cells[2], _intel(indicator))
        document.add_paragraph()

    yara = report.get("yara") or []
    if yara:
        _section_heading(document, "Matched malware signatures")
        table = _table(document, ["Signature", "Description"])
        for hit in yara:
            cells = table.add_row().cells
            _cell(cells[0], hit.get("rule") or "", mono=True)
            _cell(cells[1], (hit.get("meta") or {}).get("description", ""))
        document.add_paragraph()

    files = report.get("files") or []
    if len(files) > 1:
        _section_heading(document, "Files contained within the sample")
        table = _table(document, ["Name", "Type", "Concealed"])
        for leaf in files:
            cells = table.add_row().cells
            _cell(cells[0], leaf.get("relative_name") or "", mono=True)
            _cell(cells[1], (leaf.get("detected_type") or "").split(",")[0])
            _cell(cells[2], "Yes" if leaf.get("likely_packed") else "No")
        document.add_paragraph()

    providers = report.get("providers") or []
    if providers:
        _section_heading(document, "Intelligence sources consulted")
        table = _table(document, ["Source", "Result"])
        for provider in providers:
            cells = table.add_row().cells
            _cell(cells[0], provider.get("provider") or "")
            _cell(cells[1], _provider_result(provider))
        document.add_paragraph()

    warnings = report.get("warnings") or []
    if warnings:
        _section_heading(document, "Notes on reading this file")
        for warning in warnings:
            document.add_paragraph(warning, style="List Bullet")
        document.add_paragraph()


def _android(document: Document, report: dict[str, Any]) -> None:
    """Android sections, for reports where the sample is an app.

    The impersonation findings come first and in full. For the fraud this tool
    exists for, "the app is not who it says it is" is the finding an officer acts
    on, and it is not something the permission list conveys.
    """
    apps = [leaf for leaf in report.get("files") or [] if leaf.get("is_apk")]
    if not apps:
        return

    for app in apps:
        _section_heading(document, "Android application")

        table = document.add_table(rows=0, cols=2)
        table.style = "Table Grid"
        for label, value in [
            ("Displayed name", app.get("app_label") or "(no name)"),
            ("Internal package", app.get("package") or "—"),
        ]:
            cells = table.add_row().cells
            _cell(cells[0], label, bold=True, width=Pt(120))
            _cell(cells[1], value, mono=label == "Internal package")
        document.add_paragraph()

        signals = app.get("signals") or []
        if signals:
            _section_heading(document, "Signs this app is not what it claims")
            for signal in signals:
                paragraph = document.add_paragraph()
                run = paragraph.add_run(signal.get("plain_language", ""))
                run.bold = True
                _emphasis(document, f"Evidence: {signal.get('detail', '')}")
            document.add_paragraph()

        _permission_table(
            document,
            "Permissions granting access to personal data",
            app.get("dangerous_permissions") or [],
        )
        _permission_table(
            document,
            "Permissions commonly abused for fraud",
            app.get("high_abuse_permissions") or [],
        )

        certificates = app.get("certificates") or []
        if certificates:
            _section_heading(document, "Signing certificate")
            table = _table(document, ["Signed by", "Fingerprint (SHA-256)"])
            for certificate in certificates:
                cells = table.add_row().cells
                _cell(cells[0], certificate.get("subject", ""))
                _cell(cells[1], certificate.get("sha256", ""), mono=True)
            _emphasis(
                document,
                "The fingerprint identifies the signing key. Apps from the same "
                "campaign are usually signed with the same key even when their "
                "names and package names differ, so it is worth comparing against "
                "other exhibits.",
            )
            document.add_paragraph()


def _executed_runs(report: dict[str, Any]) -> list[dict]:
    """Detonations in which the sample actually ran.

    A failed run is deliberately excluded: nothing was observed in it, so it
    grants no licence to describe anything as observed.
    """
    return [
        run
        for run in report.get("detonations") or []
        if run.get("status") in {"complete", "timeout"}
    ]


def _observed_behaviour(document: Document, report: dict[str, Any]) -> None:
    """The timeline: what the sample was seen doing, in the order it did it.

    Kept in its own section, under its own heading, and never merged with the
    capability table. The two answer different questions and carry very
    different evidential weight.
    """
    runs = [run for run in _executed_runs(report) if run.get("events")]
    if not runs:
        return

    for run in runs:
        _section_heading(document, "Observed behaviour")

        timed = run.get("timed", True)
        _emphasis(
            document,
            "Recorded while the file was running in the sandbox. Times are "
            "measured from the moment it was started."
            if timed
            else "Recorded while the file was running in the emulator, which "
            "reports the order of events but not the time between them. The "
            "entries below are listed in order, not to a clock.",
        )

        table = _table(
            document, ["When" if timed else "Step", "What the file did", "Detail"]
        )
        for position, event in enumerate(run.get("events") or [], start=1):
            cells = table.add_row().cells
            _cell(
                cells[0],
                _offset(event.get("offset_ms")) if timed else str(position),
                mono=True,
            )
            _cell(cells[1], _describe_event(event))
            _cell(cells[2], _event_detail(event))

        if run.get("coverage"):
            _emphasis(document, run["coverage"])
        document.add_paragraph()


def _describe_event(event: dict[str, Any]) -> str:
    """One event as a sentence, prefixed so it cannot be read as a capability."""
    action = event.get("action") or "did"
    target = event.get("target") or ""
    return f"Observed: {action} {target}".strip()


def _event_detail(event: dict[str, Any]) -> str:
    parts = [event.get("detail") or ""]
    size = event.get("size_bytes")
    if size:
        parts.append(f"{size:,} bytes")
    return "; ".join(part for part in parts if part)


def _offset(offset_ms) -> str:
    if not isinstance(offset_ms, (int, float)):
        return "—"
    return f"+{offset_ms / 1000:.1f}s"


def _data_sent_out(document: Document, report: dict[str, Any]) -> None:
    """Data read from the phone and then transmitted.

    This is the finding an investigator is looking for, so it gets its own
    heading rather than being left for the reader to assemble out of the
    timeline. The wording states the pairing and the interval and stops there —
    that those exact bytes were those exact records is not something the
    observation establishes, and the report must not imply it.
    """
    findings = [
        finding
        for run in _executed_runs(report)
        for finding in run.get("exfiltration") or []
    ]
    if not findings:
        return

    _section_heading(document, "Data sent out of the device")

    table = _table(
        document, ["What was read", "Where it was sent", "Delay", "Amount sent"]
    )
    for finding in findings:
        cells = table.add_row().cells
        _cell(cells[0], finding.get("what") or "")
        _cell(cells[1], finding.get("where") or "", mono=True)
        _cell(cells[2], _gap(finding.get("gap_ms")))
        _cell(
            cells[3],
            f"{finding['bytes_sent']:,} bytes" if finding.get("bytes_sent") else "—",
        )

    _emphasis(
        document,
        "Each row records data being read from the device and a transmission "
        "leaving it shortly afterwards. The two were observed in that order and "
        "within the interval shown. Establishing that the transmission carried "
        "exactly the data that was read would require the contents of the "
        "traffic, which encryption may prevent.",
    )
    document.add_paragraph()


def _gap(gap_ms) -> str:
    if not isinstance(gap_ms, (int, float)):
        return "not measured"
    return f"{gap_ms / 1000:.1f} seconds later"


def _permission_table(document: Document, title: str, permissions: list[str]) -> None:
    if not permissions:
        return
    _section_heading(document, title)
    for permission in permissions:
        # The android.permission. prefix is on every entry and carries no
        # information for the reader.
        document.add_paragraph(
            permission.replace("android.permission.", ""), style="List Bullet"
        )
    document.add_paragraph()


def _closing(
    document: Document,
    report: dict[str, Any],
    narrative: NarrativeResult | None,
) -> None:
    _section_heading(document, "Statement of method")

    _method_statement(document, report)
    if (report.get("verdict") or "unknown") == "unknown":
        _emphasis(document, NOT_PROOF_OF_SAFETY)

    # Whether a language model touched this document is disclosed on the
    # document, not left to the reader to discover.
    if narrative and narrative.status == "ok":
        _emphasis(
            document,
            "The descriptive sections of this report were drafted by an automated "
            f"language model ({narrative.model}) from the findings listed above. "
            "The assessment, the basis for it, and all evidence tables are produced "
            "by automated analysis and are not model-generated. The narrative "
            "should be checked by the reporting officer before the report is "
            "relied upon.",
        )


def _method_statement(document: Document, report: dict[str, Any]) -> None:
    """Whether the file was run, stated plainly and only where it is true.

    Most reports remain static-only and keep the original wording. A report that
    did detonate the sample must say so, say where, and say when — an officer
    cannot assess an observation without knowing the conditions it was made in.
    """
    executed = _executed_runs(report)

    if not executed:
        _emphasis(document, NOT_EXECUTED)
        for run in report.get("detonations") or []:
            if run.get("error"):
                _emphasis(document, DETONATION_ATTEMPTED.format(error=run["error"]))
        return

    for run in executed:
        _emphasis(
            document,
            EXECUTED.format(
                when=_stamp(run.get("started_at")),
                containment=CONTAINMENT.get(
                    run.get("platform", ""), "an isolated environment."
                ),
            ),
        )


# Helpers --------------------------------------------------------------------


def _section_heading(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(4)
    _rule(paragraph)


def _table(document: Document, headers: list[str]):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for cell, header in zip(table.rows[0].cells, headers):
        _cell(cell, header, bold=True)
    return table


def _cell(cell, text: str, bold: bool = False, mono: bool = False, width=None) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(9.5)
    if mono:
        run.font.name = "Consolas"
    if width is not None:
        cell.width = width


def _emphasis(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)


def _rule(paragraph) -> None:
    """Draw a hairline under a paragraph — python-docx has no direct API."""
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "999999")
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _footer(document: Document, report: dict[str, Any]) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        f"SHA-256 {(report.get('sha256') or '')[:16]}…  ·  Page "
    )
    run.font.size = Pt(8)
    _page_number(paragraph)


def _page_number(paragraph) -> None:
    """Insert a live PAGE field so pagination is correct after editing."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run.font.size = Pt(8)
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _intel(indicator: dict[str, Any]) -> str:
    parts: list[str] = []
    threatfox = indicator.get("threatfox") or {}
    if threatfox.get("malware"):
        parts.append(
            f"Recorded in ThreatFox as infrastructure for {threatfox['malware']}"
        )
    abuseipdb = indicator.get("abuseipdb") or {}
    if abuseipdb.get("abuse_confidence") is not None:
        location = f" ({abuseipdb['country']})" if abuseipdb.get("country") else ""
        parts.append(
            f"Reported for abuse with {abuseipdb['abuse_confidence']}% "
            f"confidence{location}"
        )
    urlscan = indicator.get("urlscan") or {}
    if urlscan.get("result_count"):
        parts.append(f"{urlscan['result_count']} prior urlscan.io record(s)")
    return "; ".join(parts) if parts else "No record found"


def _provider_result(provider: dict[str, Any]) -> str:
    labels = {
        "ok": "Answered",
        "not_found": "No record of this file or indicator",
        "unavailable": "Could not be reached",
        "skipped": "Not consulted",
    }
    label = labels.get(provider.get("status", ""), provider.get("status", ""))
    detail = provider.get("detail") or ""
    return f"{label} — {detail}" if detail else label


def _stamp(iso: str | None) -> str:
    if not iso:
        return "—"
    try:
        return datetime.fromisoformat(iso).strftime("%d %B %Y, %H:%M UTC")
    except ValueError:
        return iso
