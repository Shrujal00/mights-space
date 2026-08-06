import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app.config import Settings
from app.main import create_app

FIXTURE = Path(__file__).parent / "fixtures" / "benign_pe32.exe"


@pytest.fixture
def client(tmp_path):
    rules = tmp_path / "rules"
    rules.mkdir()
    settings = Settings(
        database_url=f"sqlite:///{tmp_path / 'test.db'}",
        sample_storage_dir=tmp_path / "samples",
        yara_rules_dir=rules,
        offline_mode=True,
    )
    with TestClient(create_app(settings)) as test_client:
        yield test_client


def upload(client, name=None, content=None):
    if content is None:
        content = FIXTURE.read_bytes()
    return client.post(
        "/api/samples",
        files={"file": (name or FIXTURE.name, io.BytesIO(content), "application/octet-stream")},
    )


def test_uploading_a_file_reports_its_hash(client):
    response = upload(client)

    assert response.status_code == 201
    assert len(response.json()["sha256"]) == 64


def test_analysis_completes_and_produces_a_verdict(client):
    sample_id = upload(client).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()

    assert report["status"] == "complete"
    assert report["verdict"] in {"malicious", "suspicious", "unknown"}


def test_the_report_carries_a_plain_language_narrative(client):
    sample_id = upload(client).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()

    assert "was not run" in report["narrative"].lower()


def test_the_report_lists_the_files_found_inside_an_archive(client):
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("payload.exe", FIXTURE.read_bytes())
        archive.writestr("notes.txt", b"contact http://panel.evil.tk/gate.php")

    sample_id = upload(client, "bundle.zip", buffer.getvalue()).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()
    assert sorted(f["relative_name"] for f in report["files"]) == [
        "notes.txt",
        "payload.exe",
    ]


def test_indicators_extracted_from_the_sample_are_reported(client):
    content = b"beacon to http://panel.evil.tk/gate.php right now"
    sample_id = upload(client, "beacon.txt", content).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()

    assert "panel.evil.tk" in [i["value"] for i in report["indicators"]]


def test_capabilities_are_reported_for_a_windows_binary(client):
    sample_id = upload(client).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()

    assert isinstance(report["techniques"], list)


def test_uploading_the_same_file_twice_reuses_the_first_analysis(client):
    first = upload(client).json()
    second = upload(client).json()

    assert first["id"] == second["id"]


def test_samples_can_be_listed(client):
    upload(client)

    listing = client.get("/api/samples").json()

    assert len(listing) == 1
    assert listing[0]["filename"] == FIXTURE.name


def test_an_unknown_sample_id_returns_404(client):
    assert client.get("/api/samples/9999").status_code == 404


def test_offline_mode_marks_every_provider_as_skipped(client):
    sample_id = upload(client).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()

    assert report["providers"]
    assert all(p["status"] == "skipped" for p in report["providers"])


def test_extraction_warnings_reach_the_report(client):
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("../escaped.txt", b"escaped")

    sample_id = upload(client, "slip.zip", buffer.getvalue()).json()["id"]

    report = client.get(f"/api/samples/{sample_id}").json()
    assert any("traversal" in w.lower() for w in report["warnings"])


class TestExports:
    def test_csv_export_is_downloadable(self, client):
        content = b"beacon to http://panel.evil.tk/gate.php"
        sample_id = upload(client, "beacon.txt", content).json()["id"]

        response = client.get(f"/api/samples/{sample_id}/export.csv")

        assert response.status_code == 200
        assert "indicator_value" in response.text
        assert "panel.evil.tk" in response.text

    def test_stix_export_is_downloadable(self, client):
        content = b"beacon to http://panel.evil.tk/gate.php"
        sample_id = upload(client, "beacon.txt", content).json()["id"]

        response = client.get(f"/api/samples/{sample_id}/export.stix")

        assert response.status_code == 200
        assert response.json()["type"] == "bundle"

    def test_exporting_an_unknown_sample_returns_404(self, client):
        assert client.get("/api/samples/9999/export.csv").status_code == 404

    def test_word_report_is_downloadable(self, client):
        sample_id = upload(client).json()["id"]

        response = client.get(f"/api/samples/{sample_id}/export.docx")

        assert response.status_code == 200
        assert response.headers["content-type"].endswith("wordprocessingml.document")
        assert response.content[:2] == b"PK"  # a .docx is a zip

    def test_the_word_report_carries_the_authority_heading(self, client):
        sample_id = upload(client).json()["id"]

        response = client.get(f"/api/samples/{sample_id}/export.docx")

        document = Document(io.BytesIO(response.content))
        assert any("CYBER POLICE" in p.text for p in document.paragraphs)

    def test_offline_mode_still_produces_a_word_report_without_a_narrative(self, client):
        # Air-gapped operation cannot reach the writing model, and must not lose
        # the filed report because of it.
        sample_id = upload(client).json()["id"]

        response = client.get(f"/api/samples/{sample_id}/export.docx")

        assert response.headers["x-narrative-status"] == "skipped"
        assert response.status_code == 200

    def test_exporting_an_unknown_sample_as_word_returns_404(self, client):
        assert client.get("/api/samples/9999/export.docx").status_code == 404


def test_an_empty_upload_is_rejected(client):
    response = upload(client, "empty.bin", b"")

    assert response.status_code == 400


class TestDetonationInTheReport:
    def test_a_static_only_report_carries_an_empty_detonation_list(self, client):
        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()

        assert report["detonations"] == []

    def test_dynamic_analysis_is_off_unless_it_is_switched_on(self, client):
        """Booting an emulator takes minutes and several gigabytes. Uploading a
        file must not start one unless the operator asked for it."""
        assert client.app.state.settings.dynamic_analysis_enabled is False

    def _record_run(self, client, sample_id, **overrides):
        from datetime import datetime, timezone

        from app.analysis.behavior import BehaviorEvent, DetonationResult
        from app.detonation import record_run

        start = datetime(2026, 8, 6, 14, 3, 20, tzinfo=timezone.utc)
        result = DetonationResult(
            platform="android",
            engine="frida",
            status=overrides.get("status", "complete"),
            started_at=start,
            finished_at=start,
            coverage="The app ran for 45 seconds.",
            events=[
                BehaviorEvent.since(
                    start, start, category="data-access", action="read",
                    target="SMS inbox", detail="247 records", source="frida",
                ),
                BehaviorEvent.since(
                    start,
                    start.replace(second=23),
                    category="network", action="sent", target="185.244.25.14:443",
                    source="frida", size_bytes=34816,
                ),
            ],
        )
        database = client.app.state.database
        with database.session() as session:
            record_run(session, sample_id, result)
            session.commit()

    def test_a_recorded_run_appears_in_the_report_with_its_events(self, client):
        sample_id = upload(client).json()["id"]
        self._record_run(client, sample_id)

        report = client.get(f"/api/samples/{sample_id}").json()

        run = report["detonations"][0]
        assert run["platform"] == "android"
        assert [event["target"] for event in run["events"]] == [
            "SMS inbox",
            "185.244.25.14:443",
        ]

    def test_the_report_pairs_the_read_with_the_transmission_that_followed(self, client):
        sample_id = upload(client).json()["id"]
        self._record_run(client, sample_id)

        report = client.get(f"/api/samples/{sample_id}").json()

        exfiltration = report["detonations"][0]["exfiltration"]
        assert exfiltration[0]["where"] == "185.244.25.14:443"
        assert exfiltration[0]["bytes_sent"] == 34816

    def test_a_failed_run_reports_no_observations(self, client):
        sample_id = upload(client).json()["id"]
        self._record_run(client, sample_id, status="failed")

        report = client.get(f"/api/samples/{sample_id}").json()

        assert report["detonations"][0]["status"] == "failed"
        assert report["detonations"][0]["exfiltration"] == []

    def test_the_word_report_of_a_detonated_sample_shows_the_timeline(self, client):
        sample_id = upload(client).json()["id"]
        self._record_run(client, sample_id)

        response = client.get(f"/api/samples/{sample_id}/export.docx")
        document = Document(io.BytesIO(response.content))
        text = "\n".join(
            [p.text for p in document.paragraphs]
            + [c.text for t in document.tables for r in t.rows for c in r.cells]
        )

        assert "OBSERVED BEHAVIOUR" in text
        assert "SMS inbox" in text


class TestDetonationRuns:
    """The sandbox is stubbed here. Booting a real emulator is verified
    separately; what matters at this level is that a detonation is triggered,
    folded into the verdict, and never left the sample stuck."""

    def _stub(self, client, outcome):
        class Stub:
            calls = 0

            def engine_for(self, result):
                return "android", "frida"

            def applies(self, result):
                return True

            def can_run(self, result):
                return True

            def run(self, path, result, on_progress=None):
                Stub.calls += 1
                return outcome

        stub = Stub()
        client.app.state.detonator = stub
        return Stub

    def _outcome(self, **overrides):
        from datetime import datetime, timezone

        from app.analysis.behavior import BehaviorEvent, DetonationResult

        start = datetime(2026, 8, 6, 14, 3, 20, tzinfo=timezone.utc)
        base = {
            "platform": "android",
            "engine": "frida",
            "status": "complete",
            "started_at": start,
            "finished_at": start,
            "events": [
                BehaviorEvent.since(
                    start, start, category="data-access", action="read",
                    target="SMS inbox", detail="247 records", source="frida",
                ),
                BehaviorEvent.since(
                    start, start.replace(second=23), category="network",
                    action="sent", target="185.244.25.14:443", source="frida",
                    size_bytes=34816,
                ),
            ],
        }
        base.update(overrides)
        return DetonationResult(**base)

    def test_an_uploaded_sample_is_detonated(self, client):
        stub = self._stub(client, self._outcome())

        upload(client)

        assert stub.calls == 1

    def test_the_sample_ends_up_complete_not_stuck_detonating(self, client):
        self._stub(client, self._outcome())

        sample_id = upload(client).json()["id"]

        assert client.get(f"/api/samples/{sample_id}").json()["status"] == "complete"

    def test_observed_exfiltration_reaches_the_verdict(self, client):
        self._stub(client, self._outcome())

        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()
        assert report["verdict"] in {"suspicious", "malicious"}
        assert any("was observed reading" in reason for reason in report["reasons"])

    def test_a_detonated_report_no_longer_claims_the_file_was_never_run(self, client):
        self._stub(client, self._outcome())

        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()
        assert "was not run" not in report["narrative"].lower()

    def test_a_failed_detonation_leaves_the_static_report_intact(self, client):
        """A sandbox that will not start costs the report its dynamic section
        and nothing else."""
        from app.analysis.behavior import DetonationResult
        from datetime import datetime, timezone

        self._stub(
            client,
            DetonationResult.failed(
                platform="android",
                engine="frida",
                started_at=datetime(2026, 8, 6, tzinfo=timezone.utc),
                error="the emulator did not boot",
            ),
        )

        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()
        assert report["status"] == "complete"
        assert report["verdict"] is not None
        assert "was not run" in report["narrative"].lower()
        assert report["detonations"][0]["error"] == "the emulator did not boot"

    def test_a_sandbox_that_crashes_does_not_fail_the_sample(self, client):
        class Exploding:
            def engine_for(self, result):
                return "android", "frida"

            def applies(self, result):
                return True

            def can_run(self, result):
                return True

            def run(self, path, result, on_progress=None):
                raise RuntimeError("the sandbox fell over")

        client.app.state.detonator = Exploding()

        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()
        assert report["status"] == "complete"
        assert report["verdict"] is not None

    def test_a_stale_model_narrative_is_discarded_when_behaviour_is_added(self, client):
        """The cached prose described a static-only report. Once the sample has
        been run it is describing a different document."""
        self._stub(client, self._outcome())

        sample_id = upload(client).json()["id"]

        from app.models import Sample

        with client.app.state.database.session() as session:
            assert session.get(Sample, sample_id).narrative_sections is None


class TestListingCarriesFileType:
    def test_the_list_reports_what_each_file_was_found_to_be(self, client):
        """The interface separates Android from Windows work. Doing that on the
        filename would be defeated by a renamed .apk, which is exactly the kind
        of file this tool exists to catch, so the list carries the detected type
        the analysis established."""
        upload(client)

        listed = client.get("/api/samples").json()

        assert "PE32" in listed[0]["detected_type"]

    def test_a_file_still_being_read_lists_without_a_type_rather_than_a_wrong_one(
        self, client
    ):
        sample_id = upload(client).json()["id"]
        from app.models import Sample

        with client.app.state.database.session() as session:
            sample = session.get(Sample, sample_id)
            sample.detected_type = ""
            sample.status = "queued"
            session.commit()

        listed = client.get("/api/samples").json()

        assert listed[0]["detected_type"] == ""


class TestStartingBehaviouralAnalysis:
    """Static analysis runs on upload. Running the file is a separate, deliberate
    step the investigator asks for — it executes the sample, takes minutes, and
    only one can run at a time."""

    def _stub(self, client, *, lines=("Starting the sandbox", "Watching the app")):
        from datetime import datetime, timezone

        from app.analysis.behavior import BehaviorEvent, DetonationResult

        start = datetime(2026, 8, 6, 14, 3, 20, tzinfo=timezone.utc)

        class Stub:
            def engine_for(self, result):
                return "android", "frida"

            def applies(self, result):
                return False  # never automatic

            def can_run(self, result):
                return True

            def run(self, path, result, on_progress=None):
                for line in lines:
                    if on_progress:
                        on_progress(line)
                return DetonationResult(
                    platform="android",
                    engine="frida",
                    status="complete",
                    started_at=start,
                    finished_at=start,
                    coverage="Ran for 40 seconds.",
                    events=[
                        BehaviorEvent.since(
                            start, start, category="data-access", action="read",
                            target="SMS inbox", detail="247 records",
                            source="frida", record_count=247,
                        ),
                    ],
                )

        client.app.state.detonator = Stub()

    def test_a_completed_sample_can_be_sent_to_the_sandbox(self, client):
        self._stub(client)
        sample_id = upload(client).json()["id"]

        response = client.post(f"/api/samples/{sample_id}/detonate")

        assert response.status_code == 202

    def test_the_run_produces_observations_on_the_report(self, client):
        self._stub(client)
        sample_id = upload(client).json()["id"]

        client.post(f"/api/samples/{sample_id}/detonate")

        report = client.get(f"/api/samples/{sample_id}").json()
        assert report["detonations"][0]["events"][0]["target"] == "SMS inbox"

    def test_progress_is_recorded_so_the_wait_is_not_silent(self, client):
        """A detonation takes minutes. Someone watching needs to see that
        something is happening, and it has to be what is actually happening."""
        self._stub(client, lines=("Starting the sandbox", "Installing the app"))
        sample_id = upload(client).json()["id"]

        client.post(f"/api/samples/{sample_id}/detonate")

        run = client.get(f"/api/samples/{sample_id}").json()["detonations"][0]
        assert [entry["message"] for entry in run["progress"]] == [
            "Starting the sandbox",
            "Installing the app",
        ]

    def test_each_progress_entry_is_stamped_with_when_it_happened(self, client):
        self._stub(client)
        sample_id = upload(client).json()["id"]

        client.post(f"/api/samples/{sample_id}/detonate")

        run = client.get(f"/api/samples/{sample_id}").json()["detonations"][0]
        assert run["progress"][0]["at"] is not None

    def test_a_sample_that_has_not_finished_reading_cannot_be_detonated(self, client):
        self._stub(client)
        sample_id = upload(client).json()["id"]
        from app.models import Sample

        with client.app.state.database.session() as session:
            session.get(Sample, sample_id).status = "queued"
            session.commit()

        assert client.post(f"/api/samples/{sample_id}/detonate").status_code == 409

    def test_an_unknown_sample_cannot_be_detonated(self, client):
        self._stub(client)

        assert client.post("/api/samples/9999/detonate").status_code == 404

    def test_the_sample_is_left_complete_and_not_stuck_running(self, client):
        self._stub(client)
        sample_id = upload(client).json()["id"]

        client.post(f"/api/samples/{sample_id}/detonate")

        assert client.get(f"/api/samples/{sample_id}").json()["status"] == "complete"

    def test_a_sandbox_that_is_not_configured_says_so_rather_than_failing_quietly(
        self, client
    ):
        class Unconfigured:
            def engine_for(self, result):
                return "android", "frida"

            def applies(self, result):
                return False

            def can_run(self, result):
                return False

            def run(self, path, result, on_progress=None):
                raise AssertionError("must not be called")

        client.app.state.detonator = Unconfigured()
        sample_id = upload(client).json()["id"]

        response = client.post(f"/api/samples/{sample_id}/detonate")

        assert response.status_code == 409
        assert "sandbox" in response.json()["detail"].lower()

    def test_the_report_says_whether_the_file_can_be_sent_to_the_sandbox(self, client):
        """The interface needs to know whether to offer the button at all,
        without guessing at the server's configuration."""
        self._stub(client)
        sample_id = upload(client).json()["id"]

        report = client.get(f"/api/samples/{sample_id}").json()

        assert report["can_detonate"] is True
