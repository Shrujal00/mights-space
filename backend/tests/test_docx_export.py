import zipfile
from io import BytesIO

from docx import Document

from app.analysis.export.docx_export import report_to_docx
from app.analysis.narrative import NarrativeResult

AUTHORITY = "SURAT CYBER POLICE INDIA"


def report(**overrides):
    base = {
        "id": 1,
        "filename": "invoice.exe",
        "detected_type": "PE32 executable (GUI) Intel 80386",
        "size": 486912,
        "sha256": "a" * 64,
        "sha1": "b" * 40,
        "md5": "c" * 32,
        "verdict": "malicious",
        "headline": "This file is malicious.",
        "narrative": "invoice.exe\n\nThis file is malicious.\n\nThe file was not run.",
        "reasons": ["45 of 70 antivirus engines identify this file as malicious."],
        "techniques": [
            {
                "technique_id": "T1056.001",
                "name": "Input Capture: Keylogging",
                "plain_language": "Can record everything typed on the keyboard.",
                "evidence": ["SetWindowsHookExA"],
                "basis": "static-import",
            }
        ],
        "indicators": [
            {
                "type": "ipv4",
                "value": "185.244.25.14",
                "threatfox": {"malware": "AgentTesla", "threat_type": "botnet_cc"},
                "abuseipdb": {"abuse_confidence": 100, "country": "RU"},
                "urlscan": None,
            }
        ],
        "yara": [
            {
                "rule": "MAL_AgentTesla_Nov23",
                "namespace": "",
                "tags": [],
                "meta": {"description": "Detects AgentTesla"},
            }
        ],
        "files": [],
        "providers": [{"provider": "virustotal", "status": "ok", "detail": ""}],
        "warnings": [],
        "created_at": "2026-08-06T10:00:00+00:00",
        "completed_at": "2026-08-06T10:00:05+00:00",
    }
    base.update(overrides)
    return base


def written(**overrides):
    return NarrativeResult(
        status="ok",
        model="gemma4:31b",
        sections={
            "overview": "A Windows program submitted for examination.",
            "assessment": "The file is malicious.",
            "capabilities": "It is able to record what is typed.",
            "destinations": "It refers to one internet address.",
            "limitations": "Only the file's code was read.",
            **overrides,
        },
    )


def text_of(blob: bytes) -> str:
    """All prose in the document, paragraphs and tables alike."""
    document = Document(BytesIO(blob))
    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


class TestDocument:
    def test_produces_a_readable_word_file(self):
        blob = report_to_docx(report(), AUTHORITY, written())

        # A .docx is a zip; Word will not open it without the document part.
        with zipfile.ZipFile(BytesIO(blob)) as archive:
            assert "word/document.xml" in archive.namelist()

    def test_carries_the_authority_at_the_head(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert AUTHORITY in text

    def test_states_the_verdict(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "MALICIOUS" in text

    def test_records_the_hashes_that_identify_the_exhibit(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "a" * 64 in text
        assert "c" * 32 in text

    def test_lists_the_basis_for_the_assessment(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "45 of 70 antivirus engines" in text

    def test_includes_capabilities_and_the_evidence_for_them(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "record everything typed" in text
        assert "SetWindowsHookExA" in text

    def test_includes_network_destinations_with_their_intelligence(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "185.244.25.14" in text
        assert "AgentTesla" in text


class TestHonesty:
    def test_a_static_only_report_states_the_file_was_not_run(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "was not run" in text

    def test_an_unknown_verdict_says_it_is_not_proof_of_safety(self):
        # The document may be acted on. Absence of findings is not a clearance.
        text = text_of(
            report_to_docx(
                report(verdict="unknown", headline="No known indicators were found."),
                AUTHORITY,
                written(),
            )
        )

        assert "not proof that the file is harmless" in text

    def test_a_model_written_narrative_is_disclosed_in_the_document(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "language model" in text
        assert "gemma4:31b" in text

    def test_nothing_is_attributed_to_a_model_when_none_was_used(self):
        text = text_of(
            report_to_docx(
                report(), AUTHORITY, NarrativeResult(status="skipped")
            )
        )

        assert "language model" not in text


class TestFallback:
    def test_the_deterministic_summary_is_used_when_no_narrative_was_written(self):
        text = text_of(
            report_to_docx(
                report(), AUTHORITY, NarrativeResult(status="unavailable")
            )
        )

        assert "This file is malicious." in text

    def test_a_report_is_still_produced_with_no_narrative_at_all(self):
        blob = report_to_docx(report(), AUTHORITY, None)

        assert text_of(blob).count(AUTHORITY) == 1

    def test_the_evidence_tables_do_not_depend_on_the_model(self):
        # The model drafts prose only. If it is unreachable the findings must
        # still reach the document in full.
        text = text_of(
            report_to_docx(report(), AUTHORITY, NarrativeResult(status="skipped"))
        )

        assert "185.244.25.14" in text
        assert "MAL_AgentTesla_Nov23" in text
        assert "45 of 70 antivirus engines" in text


APK_LEAF = {
    "relative_name": "SBI YONO REWARDZ.apk",
    "sha256": "d" * 64,
    "detected_type": "Android package",
    "size": 4096,
    "is_pe": False,
    "likely_packed": False,
    "packing_reasons": [],
    "imported_dlls": [],
    "sections": [],
    "is_apk": True,
    "package": "com.facebook.smsrecevies",
    "app_label": "YONO SBI",
    "permissions": ["android.permission.READ_SMS"],
    "dangerous_permissions": ["android.permission.READ_SMS"],
    "high_abuse_permissions": ["android.permission.SYSTEM_ALERT_WINDOW"],
    "components": {"activities": [], "services": [], "receivers": [], "providers": []},
    "certificates": [
        {"subject": "Common Name: Android", "issuer": "Common Name: Android",
         "sha256": "e" * 64, "self_signed": True}
    ],
    "signals": [
        {
            "code": "brand-impersonation",
            "plain_language": "The app presents itself as YONO SBI but was not published by them.",
            "detail": "package com.facebook.smsrecevies does not match",
        }
    ],
    "notable_strings": [],
}


class TestAndroidReport:
    def test_names_both_what_the_app_claims_and_what_it_is(self):
        text = text_of(report_to_docx(report(files=[APK_LEAF]), AUTHORITY, written()))

        assert "YONO SBI" in text
        assert "com.facebook.smsrecevies" in text

    def test_states_the_impersonation_finding_in_plain_words(self):
        text = text_of(report_to_docx(report(files=[APK_LEAF]), AUTHORITY, written()))

        assert "was not published by them" in text

    def test_lists_permissions_without_the_android_prefix(self):
        text = text_of(report_to_docx(report(files=[APK_LEAF]), AUTHORITY, written()))

        assert "READ_SMS" in text
        assert "android.permission.READ_SMS" not in text

    def test_records_the_signing_fingerprint_for_linking_exhibits(self):
        # Apps from one campaign share a signing key even when everything else
        # about them differs.
        text = text_of(report_to_docx(report(files=[APK_LEAF]), AUTHORITY, written()))

        assert "e" * 64 in text

    def test_a_windows_report_has_no_android_section(self):
        text = text_of(report_to_docx(report(), AUTHORITY, written()))

        assert "ANDROID APPLICATION" not in text


def detonation(**overrides):
    base = {
        "platform": "android",
        "engine": "frida",
        "status": "complete",
        "started_at": "2026-08-06T14:03:20+00:00",
        "finished_at": "2026-08-06T14:05:20+00:00",
        "error": "",
        "timed": True,
        "coverage": "The app ran under instrumentation for 45 seconds.",
        "events": [
            {
                "at": "2026-08-06T14:03:22+00:00",
                "offset_ms": 2000,
                "category": "data-access",
                "action": "read",
                "target": "SMS inbox",
                "detail": "247 records",
                "source": "frida",
                "size_bytes": None,
            },
            {
                "at": "2026-08-06T14:03:25+00:00",
                "offset_ms": 5000,
                "category": "network",
                "action": "sent",
                "target": "185.244.25.14:443",
                "detail": "",
                "source": "frida",
                "size_bytes": 34816,
            },
        ],
        "exfiltration": [
            {
                "what": "247 records from SMS inbox",
                "where": "185.244.25.14:443",
                "when": "2026-08-06T14:03:22+00:00",
                "gap_ms": 3000,
                "bytes_sent": 34816,
                "confidence": "strong",
            }
        ],
    }
    base.update(overrides)
    return base


class TestObservedBehaviour:
    def test_a_detonated_report_has_a_timeline_of_what_happened(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "OBSERVED BEHAVIOUR" in text
        assert "SMS inbox" in text
        assert "247 records" in text

    def test_the_timeline_shows_when_each_thing_happened(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "+2.0s" in text

    def test_an_untimed_run_numbers_its_events_instead_of_timing_them(self):
        """The Windows emulator records the order of calls, not the clock. A
        timeline showing "+0.0s" against every line would imply measurements
        that were never taken."""
        text = text_of(
            report_to_docx(
                report(
                    detonations=[
                        detonation(platform="windows", engine="speakeasy", timed=False)
                    ]
                ),
                AUTHORITY,
                written(),
            )
        )

        assert "+2.0s" not in text
        assert "in order" in text.lower()

    def test_observed_behaviour_is_labelled_as_observed_not_as_capability(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "Observed" in text

    def test_how_far_the_run_got_is_stated(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "45 seconds" in text


class TestDataLeavingTheDevice:
    def test_the_report_has_a_section_for_data_sent_out(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "DATA SENT OUT OF THE DEVICE" in text

    def test_it_names_what_was_taken_and_where_it_went(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "247 records from SMS inbox" in text
        assert "185.244.25.14:443" in text

    def test_it_states_the_delay_between_reading_and_sending(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "3.0" in text

    def test_a_run_that_observed_nothing_leaving_has_no_such_section(self):
        text = text_of(
            report_to_docx(
                report(detonations=[detonation(exfiltration=[])]), AUTHORITY, written()
            )
        )

        assert "DATA SENT OUT OF THE DEVICE" not in text


class TestWhetherTheFileWasRun:
    def test_a_detonated_report_does_not_claim_the_file_was_never_run(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "was not run at any point" not in text

    def test_a_detonated_report_says_where_and_when_it_was_run(self):
        text = text_of(
            report_to_docx(report(detonations=[detonation()]), AUTHORITY, written())
        )

        assert "sandbox" in text.lower()
        assert "2026" in text

    def test_a_failed_detonation_still_states_the_file_was_never_run(self):
        text = text_of(
            report_to_docx(
                report(
                    detonations=[
                        detonation(status="failed", events=[], exfiltration=[],
                                   error="the emulator did not boot")
                    ]
                ),
                AUTHORITY,
                written(),
            )
        )

        assert "was not run at any point" in text

    def test_a_failed_detonation_says_it_was_attempted_and_why_it_failed(self):
        text = text_of(
            report_to_docx(
                report(
                    detonations=[
                        detonation(status="failed", events=[], exfiltration=[],
                                   error="the emulator did not boot")
                    ]
                ),
                AUTHORITY,
                written(),
            )
        )

        assert "the emulator did not boot" in text
